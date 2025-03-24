import pandas as pd
import docx
from PIL import Image
from content_processor.image_processor import check_relevance  # Import the image relevance logic

def export_table_to_docx(table: pd.DataFrame, doc: docx.Document):
    """export a pandas DataFrame to a table in a docx document."""
    table_in_docx = doc.add_table(rows=1, cols=len(table.columns))
    hdr_cells = table_in_docx.rows[0].cells
    for i, column_name in enumerate(table.columns):
        hdr_cells[i].text = str(column_name)
    for _, row in table.iterrows():
        row_cells = table_in_docx.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)

def export_images_and_text_to_docx(data_dict, output_file_path="/Users/kashishmandhane/Documents/Kashish Data/LAPTOP STUFF/DJ Sanghvi College/Extra-curriculars/Hackathons/Ed-tech/output/final_consolidated_notes.docx"):
    """generates a DOCX document with summarized text and inserts relevant images in between"""
    doc = docx.Document()
    text_list = data_dict.get("text", [])
    images = data_dict.get("images", [])
    img_index = 0  # To track inserted images
    for text in text_list:
        doc.add_paragraph(text)  # Add text
        if img_index < len(images):  # Ensure images are available
            image = Image.open(images[img_index]).convert("RGB")
            relevance_score = check_relevance(image, text)
            if relevance_score > 2.0:  # Only insert highly relevant images
                doc.add_picture(images[img_index], width=docx.shared.Inches(4))
                img_index += 1  # Move to the next image
    doc.save(output_file_path)

def export_tables_to_docx(data_dict, output_file_path="/Users/kashishmandhane/Documents/Kashish Data/LAPTOP STUFF/DJ Sanghvi College/Extra-curriculars/Hackathons/Ed-tech/output/consolidated_tables.docx"):
    """exporting all tables from input into docx document."""
    doc = docx.Document()
    for table in data_dict.get("tables", []):
        export_table_to_docx(table, doc)
    doc.save(output_file_path)