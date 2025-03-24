import fitz
import pytesseract
from PIL import Image
import io
import camelot
import pandas as pd
import os
import zipfile
from docx import Document

def extract_pdf_content(pdf_file):
    """
    Extracts text, images, and tables from a PDF document.
    Supports both text-based and scanned PDFs (via OCR).
    """
    pdf_content = {"text": [], "images": [], "tables": []}
    document = fitz.open(pdf_file)
    
    for page_num, page in enumerate(document):
        page_text = page.get_text()
        if page_text.strip():
            try:
                tables = camelot.read_pdf(pdf_file, pages=str(page_num+1), flavor="lattice")
                for table in tables:
                    df = pd.DataFrame(table.df.values)
                    df.columns = df.iloc[0]
                    df = df.drop(0).reset_index(drop=True)
                    pdf_content["tables"].append(df)
            except Exception as e:
                print(f"No tables found on page {page_num+1}: {e}")
            pdf_content["text"].append(page_text)
            
            for img_index, img in enumerate(page.get_images(full=True)):
                xref = img[0]
                base_image = document.extract_image(xref)
                img_bytes = base_image["image"]
                img_extension = base_image["ext"]
                image = Image.open(io.BytesIO(img_bytes))
                image = image.convert("RGBA") if image.mode in ("RGBA", "LA") else image.convert("RGB")
                img_filename = f"{pdf_file}_page{page_num+1}_img{img_index+1}.{img_extension}"
                image.save(img_filename, format=img_extension.upper())
                pdf_content["images"].append(img_filename)
        else:
            pix = page.get_pixmap()
            img = Image.open(io.BytesIO(pix.tobytes()))
            ocr_text = pytesseract.image_to_string(img)
            pdf_content["text"].append(ocr_text)
    
    document.close()
    return pdf_content

def extract_docx_content(docx_file):
    """
    Extracts text, tables, and images from a DOCX document.
    """
    docx_content = {"text": [], "images": [], "tables": []}
    doc = Document(docx_file)
    
    for para in doc.paragraphs:
        if para.text.strip():
            docx_content["text"].append(para.text.strip())
    
    for table in doc.tables:
        table_data = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        docx_content["tables"].append(pd.DataFrame(table_data))
    
    os.makedirs("media", exist_ok=True)
    with zipfile.ZipFile(docx_file, 'r') as docx:
        media_files = [f for f in docx.namelist() if f.startswith('word/media/') and f.lower().endswith(('.jpg', '.png', '.jpeg'))]
        for media in media_files:
            img_bytes = docx.read(media)
            img_extension = media.split('.')[-1]
            img_filename = f"media/{os.path.basename(media)}"
            with open(img_filename, 'wb') as img_file:
                img_file.write(img_bytes)
            docx_content["images"].append(img_filename)
    
    return docx_content
